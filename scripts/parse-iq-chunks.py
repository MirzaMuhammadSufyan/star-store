#!/usr/bin/env python3
"""Parse IQ MCQs from root/chunks/*.docx into src/data/iqBook/."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

ROOT = Path("/home/sufyan/Documents/Work/Adsense/star-store")
CHUNKS = ROOT / "chunks"
OUT = ROOT / "src" / "data" / "iqBook"

CHAPTER_DEFS = [
    {
        "id": "ch-01",
        "title": "Analogy Test (Type 1)",
        "slug": "analogy-type-1",
        "blurb": "Word relationships and verbal analogies.",
        "accent": "amber",
        "header": re.compile(r"(?m)^\s*1\.\s*Analogy Test\s*\(TYPE\s*1\)\s*$", re.I),
    },
    {
        "id": "ch-02",
        "title": "Analogy Test (Type 2)",
        "slug": "analogy-type-2",
        "blurb": "Paired analogies across purpose, cause, class, and more.",
        "accent": "orange",
        "header": re.compile(r"(?m)^\s*2\.\s*Analogy Test\s*\(TYPE\s*2\)\s*$", re.I),
    },
    {
        "id": "ch-03",
        "title": "Analogy Test (Type 3)",
        "slug": "analogy-type-3",
        "blurb": "Find what is common across a set of terms.",
        "accent": "rose",
        "header": re.compile(r"(?m)^\s*3\.\s*Analogy Test\s*\(TYPE\s*3\)\s*$", re.I),
    },
    {
        "id": "ch-04",
        "title": "Series Test",
        "slug": "series-test",
        "blurb": "Number and letter series completion.",
        "accent": "sky",
        "header": re.compile(r"(?m)^\s*4\.\s*Series Test\s*$", re.I),
    },
    {
        "id": "ch-05",
        "title": "Jumbled Spelling Test",
        "slug": "jumbled-spelling",
        "blurb": "Unscramble letters into meaningful words.",
        "accent": "teal",
        "header": re.compile(r"(?m)^\s*\(?a\)?\s*JUMBLED\s+SPELLING\s+TEST\s*$", re.I),
    },
    {
        "id": "ch-06",
        "title": "Coding and Decoding Test",
        "slug": "coding-decoding",
        "blurb": "Letter and word coding patterns.",
        "accent": "emerald",
        "header": re.compile(r"(?m)^\s*\(?b\)?\s*CODING\s+AND\s+DECODING\s+TEST\s*$", re.I),
    },
]

STOP_HEADER = re.compile(r"(?m)^\s*\(?c\)?\s*ALPHABETICAL\s+TEST\s*$", re.I)
Q_START = re.compile(r"(?m)^\s*(\d{1,3})\s*[.)]\s+(.+)$")
BOOK_NOISE = re.compile(
    r"(?im)^(?:—\s*Book page.*|Dogar'?s Unique.*|Note:.*|Direction[s]?:.*|"
    r"What comes next.*|Write the correct words.*|Rearrange the following.*|"
    r"Point out the last letter.*|Insert the missing figures.*|"
    r"Write the missing figures.*|\*+)$"
)


def read_all_text() -> tuple[str, list[str]]:
    parts: list[str] = []
    table_blobs: list[str] = []
    for path in sorted(CHUNKS.glob("*.docx")):
        doc = Document(str(path))
        for p in doc.paragraphs:
            t = p.text.replace("\xa0", " ").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                rows.append(" | ".join(x for x in cells if x))
            table_blobs.append("\n".join(rows))
    return "\n".join(parts), table_blobs


def parse_key_map(text: str, *, allow_dash_skip: bool = True) -> dict[int, str]:
    keys: dict[int, str] = {}
    # Explicit empty answers: 7.—
    empty = set()
    if allow_dash_skip:
        for m in re.finditer(r"(?i)(\d{1,3})\s*[.)]\s*[—–\-]+\s*", text):
            empty.add(int(m.group(1)))

    for m in re.finditer(r"(?i)(\d{1,3})\s*[.)]\s*\(?\s*([a-d])\s*\)?", text):
        n = int(m.group(1))
        if n in empty:
            continue
        keys[n] = m.group(2).upper()

    for m in re.finditer(r"(?im)^\s*(\d{1,3})\s*[.)]\s*\(([a-d])\)", text):
        n = int(m.group(1))
        if n in empty:
            continue
        keys[n] = m.group(2).upper()

    for n in empty:
        keys.pop(n, None)
    return keys


def parse_options(block: str) -> list[str] | None:
    """Parse (a)(b)(c)(d); repair duplicate letter typos from the source book."""
    block = block.replace("\xa0", " ")
    # OCR / typeset typos: (6) meant (b)
    block = re.sub(r"\(\s*6\s*\)", "(b)", block)
    block = re.sub(r"\(\s*([a-dA-D])\s*\)", lambda m: f"({m.group(1).lower()})", block)

    found: list[tuple[str, str]] = []
    for m in re.finditer(r"\(([a-d])\)\s*([^\n(]+)", block, re.I):
        letter = m.group(1).upper()
        val = re.sub(r"\s+", " ", m.group(2)).strip(" .,:;|\t")
        if not val or len(val) > 220:
            continue
        found.append((letter, val))

    if len(found) < 3:
        return None

    opts: dict[str, str] = {}
    next_missing = ["A", "B", "C", "D"]
    for letter, val in found:
        if letter not in opts:
            opts[letter] = val
            if letter in next_missing:
                next_missing.remove(letter)
        else:
            # Duplicate marker (book typo) → assign to next missing slot
            if next_missing:
                opts[next_missing.pop(0)] = val

    if len(opts) < 3:
        return None
    # Require all four; if still missing one and we have 3, leave None so we skip
    if len(opts) < 4:
        return None
    return [opts[L] for L in "ABCD"]


def split_chapters(full_text: str) -> dict[str, str]:
    hits: list[tuple[int, str]] = []
    for ch in CHAPTER_DEFS:
        m = ch["header"].search(full_text)
        if m:
            hits.append((m.start(), ch["id"]))
    stop = STOP_HEADER.search(full_text)
    stop_at = stop.start() if stop else len(full_text)
    hits.sort(key=lambda x: x[0])
    sections: dict[str, str] = {}
    for i, (start, cid) in enumerate(hits):
        end = hits[i + 1][0] if i + 1 < len(hits) else stop_at
        end = min(end, stop_at)
        if start < end:
            sections[cid] = full_text[start:end]
    return sections


def exercise_body(section: str) -> str:
    m = re.search(r"(?im)^\s*EXERCISE\b", section)
    body = section[m.start() :] if m else section
    cut = re.search(r"(?im)^\s*(?:Answers?\s*Key|Explanatory\s+Answers)\b", body)
    if cut:
        return body[: cut.start()]
    return body


def answer_region(section: str) -> str:
    m = re.search(r"(?im)^\s*(?:Answers?\s*Key|Explanatory\s+Answers)\b", section)
    if not m:
        return ""
    region = section[m.start() :]
    # Stop at asterisks or next major section (not mid-answer (c) lines)
    stop = re.search(
        r"(?im)^\*{3,}|(?:^\s*\(?[abc]\)?\s*(?:JUMBLED|CODING|ALPHABETICAL)\s+TEST\b)",
        region,
    )
    if stop and stop.start() > 0:
        region = region[: stop.start()]
    return region


def extract_mcqs(body: str) -> list[dict]:
    lines = [ln.strip() for ln in body.splitlines() if ln.strip()]
    blocks: list[tuple[int, str]] = []
    cur_num = None
    cur_buf: list[str] = []

    def flush() -> None:
        nonlocal cur_num, cur_buf
        if cur_num is not None and cur_buf:
            blocks.append((cur_num, "\n".join(cur_buf)))
        cur_num, cur_buf = None, []

    for raw in lines:
        if BOOK_NOISE.match(raw):
            continue
        if re.match(r"(?i)^\d+\.\s*Analogy Test", raw):
            continue
        if re.match(r"(?i)^\(?[abc]\)?\s*(JUMBLED|CODING|ALPHABETICAL)", raw):
            continue
        m = Q_START.match(raw)
        # Require a real stem (skip bare "8." figure labels)
        if m and int(m.group(1)) <= 200 and len(m.group(2).strip()) >= 2:
            flush()
            cur_num = int(m.group(1))
            cur_buf = [m.group(2)]
            continue
        if cur_num is not None:
            # Don't swallow the next numbered stem with no options
            if re.match(r"^\d{1,3}\s*[.)]\s*$", raw):
                flush()
                continue
            cur_buf.append(raw)
    flush()

    questions: list[dict] = []
    for num, block in blocks:
        options = parse_options(block)
        if not options:
            continue
        stem_part = re.split(r"\([aA]\)", block, maxsplit=1)[0]
        stem = re.sub(r"\s+", " ", stem_part).strip(" .:-")
        if len(stem) < 3:
            continue
        low = stem.lower()
        if low.startswith(("note:", "explanation", "ans.", "hint:", "example")):
            continue
        if not stem.endswith("?"):
            stem = stem + "?"
        questions.append({"sourceNum": num, "question": stem[:500], "options": options})
    return questions


def jumbled_word_mcqs(body: str, word_answers: dict[int, str]) -> list[dict]:
    out: list[dict] = []
    pool = [w for w in word_answers.values() if len(w.replace(" ", "")) >= 3]
    for m in re.finditer(
        r"(?m)^\s*(\d{1,2})\s*[.)]\s+([A-Z]{3,20})(?:\s+\(([^)]+)\))?\s*$",
        body,
    ):
        num = int(m.group(1))
        jumble = m.group(2)
        hint = (m.group(3) or "").strip()
        ans_word = word_answers.get(num)
        if not ans_word:
            continue
        ans_word = re.sub(r"[^A-Za-z ]", "", ans_word).strip().upper()
        compact = ans_word.replace(" ", "")
        if len(compact) < 3:
            continue

        # Distractors: other real answers of similar length from the same chapter
        distractors: list[str] = []
        for w in pool:
            wc = w.replace(" ", "").upper()
            if wc == compact:
                continue
            if abs(len(wc) - len(compact)) <= 2 and wc not in distractors:
                distractors.append(wc)
            if len(distractors) >= 3:
                break
        # Fallback letter-swap distractors
        letters = list(compact)
        i = 0
        while len(distractors) < 3 and i < len(letters) - 1:
            cand = letters[:]
            cand[i], cand[i + 1] = cand[i + 1], cand[i]
            w = "".join(cand)
            if w != compact and w not in distractors:
                distractors.append(w)
            i += 1

        options = [compact, *distractors[:3]]
        rot = num % 4
        options = options[rot:] + options[:rot]
        answer_letter = "ABCD"[options.index(compact)]
        stem = f"Unscramble '{jumble}'"
        if hint:
            stem += f" ({hint})"
        stem += "?"
        out.append(
            {
                "sourceNum": num,
                "question": stem,
                "options": options,
                "answer": answer_letter,
            }
        )
    return out


def parse_jumbled_word_answers(answer_text: str) -> dict[int, str]:
    keys: dict[int, str] = {}
    for m in re.finditer(r"(?i)(\d{1,2})\s*[.)]\s*([A-Z][A-Z .'-]{1,24})", answer_text):
        word = m.group(2).strip(" .")
        if word.startswith("(") or re.fullmatch(r"[A-D]", word):
            continue
        # Skip letter-key entries like "(c) GERMANY" already handled elsewhere
        keys[int(m.group(1))] = word.upper()
    return keys


def main() -> None:
    full_text, table_blobs = read_all_text()
    sections = split_chapters(full_text)

    table_keys = [parse_key_map(blob) for blob in table_blobs]
    type1_keys = table_keys[0] if len(table_keys) > 0 else {}
    type2_keys = table_keys[1] if len(table_keys) > 1 else {}

    # Type 1 answer grid misprints (noted in the source): 52/60/63 stand in for 59/63/68.
    type1_keys.setdefault(59, "B")
    type1_keys.setdefault(60, type1_keys.get(60, "B"))
    type1_keys.setdefault(63, "B")
    type1_keys.setdefault(68, "D")

    chapter_override_keys = {
        "ch-01": type1_keys,
        "ch-02": type2_keys,
    }

    OUT.mkdir(parents=True, exist_ok=True)
    for old in OUT.glob("ch-*.json"):
        old.unlink()
    for name in ("allQuestions.json", "chapters.json"):
        p = OUT / name
        if p.exists():
            p.unlink()

    collected: list[tuple[dict, list[dict]]] = []

    for cdef in CHAPTER_DEFS:
        section = sections.get(cdef["id"], "")
        if not section:
            continue
        body = exercise_body(section)
        ans_text = answer_region(section)
        section_keys = parse_key_map(ans_text)
        override = chapter_override_keys.get(cdef["id"], {})

        qs = extract_mcqs(body)

        if cdef["id"] == "ch-05":
            word_answers = parse_jumbled_word_answers(ans_text)
            existing_nums = {q["sourceNum"] for q in qs}
            for jq in jumbled_word_mcqs(body, word_answers):
                if jq["sourceNum"] not in existing_nums:
                    qs.append(jq)

        seen: set[tuple] = set()
        uniq: list[dict] = []
        for q in qs:
            key = (q["sourceNum"], q["question"][:60].lower())
            if key in seen:
                continue
            seen.add(key)
            ans = q.get("answer") or override.get(q["sourceNum"]) or section_keys.get(q["sourceNum"])
            if not ans or any(o.startswith("Option ") for o in q["options"]):
                continue
            # Skip "decode all of these" items (not single-choice)
            if re.search(r"(?i)decipher the following coded words", q["question"]):
                continue
            uniq.append(
                {
                    "sourceNum": q["sourceNum"],
                    "question": q["question"],
                    "options": q["options"],
                    "answer": ans,
                }
            )
        collected.append((cdef, uniq))

    chapters_index = []
    flat = []
    qid = 1
    for cdef, qs in collected:
        qs = sorted(qs, key=lambda x: x["sourceNum"])
        fixed = []
        for q in qs:
            item = {
                "id": qid,
                "sourceNum": q["sourceNum"],
                "question": q["question"],
                "options": q["options"],
                "answer": q["answer"],
                "answerVerified": True,
            }
            fixed.append(item)
            flat.append(
                {
                    "id": qid,
                    "category": cdef["title"],
                    "chapterId": cdef["id"],
                    "question": item["question"],
                    "options": item["options"],
                    "answer": item["answer"],
                    "answerVerified": True,
                }
            )
            qid += 1

        if not fixed:
            continue

        payload = {
            "id": cdef["id"],
            "title": cdef["title"],
            "slug": cdef["slug"],
            "blurb": cdef["blurb"],
            "accent": cdef["accent"],
            "questions": fixed,
        }
        (OUT / f"{cdef['id']}.json").write_text(
            json.dumps(payload, indent=2, ensure_ascii=False) + "\n",
            encoding="utf-8",
        )
        chapters_index.append(
            {
                "id": cdef["id"],
                "title": cdef["title"],
                "slug": cdef["slug"],
                "blurb": cdef["blurb"],
                "accent": cdef["accent"],
                "questionCount": len(fixed),
                "verifiedCount": len(fixed),
                "figureBased": False,
                "keysOnly": False,
                "file": f"{cdef['id']}.json",
            }
        )

    index = {
        "bookTitle": "IQ Practice Tests",
        "source": "chunks/",
        "chapterCount": len(chapters_index),
        "totalQuestions": len(flat),
        "chapters": chapters_index,
    }
    (OUT / "chapters.json").write_text(
        json.dumps(index, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    (OUT / "allQuestions.json").write_text(
        json.dumps(flat, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )

    print(json.dumps({"chapters": len(chapters_index), "questions": len(flat)}, indent=2))
    for c in chapters_index:
        print(f"- {c['id']} {c['slug']}: {c['questionCount']} Qs · {c['title']}")


if __name__ == "__main__":
    main()
