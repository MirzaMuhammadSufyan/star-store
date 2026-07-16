#!/usr/bin/env python3
"""Parse IQ verbal MCQs + theory from Dogar_IQ_Verbal_Part.docx into src/data/iqBook/."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

ROOT = Path("/home/sufyan/Documents/Work/Adsense/star-store")
SOURCE = ROOT / "Dogar_IQ_Verbal_Part.docx"
OUT = ROOT / "src" / "data" / "iqBook"

CHAPTER_DEFS = [
    {
        "id": "ch-01",
        "title": "Analogy Test (Type 1)",
        "slug": "analogy-type-1",
        "blurb": "Word relationships and verbal analogies.",
        "accent": "amber",
        "header": re.compile(r"(?m)^\s*1\.\s*Analogy Test\s*\(TYPE\s*1\)\s*$", re.I),
    },
    {
        "id": "ch-02",
        "title": "Analogy Test (Type 2)",
        "slug": "analogy-type-2",
        "blurb": "Paired analogies across purpose, cause, class, and more.",
        "accent": "orange",
        "header": re.compile(r"(?m)^\s*2\.\s*Analogy Test\s*\(TYPE\s*2\)\s*$", re.I),
    },
    {
        "id": "ch-03",
        "title": "Analogy Test (Type 3)",
        "slug": "analogy-type-3",
        "blurb": "Find what is common across a set of terms.",
        "accent": "rose",
        "header": re.compile(r"(?m)^\s*3\.\s*Analogy Test\s*\(TYPE\s*3\)\s*$", re.I),
    },
    {
        "id": "ch-04",
        "title": "Series Test",
        "slug": "series-test",
        "blurb": "Number and letter series completion.",
        "accent": "sky",
        "header": re.compile(r"(?m)^\s*4\.\s*Series Test\s*$", re.I),
    },
    {
        "id": "ch-05",
        "title": "Jumbled Spelling Test",
        "slug": "jumbled-spelling",
        "blurb": "Unscramble letters into meaningful words.",
        "accent": "teal",
        "header": re.compile(r"(?m)^\s*\(?a\)?\s*JUMBLED\s+SPELLING\s+TEST\s*$", re.I),
    },
    {
        "id": "ch-06",
        "title": "Coding and Decoding Test",
        "slug": "coding-decoding",
        "blurb": "Letter and word coding patterns.",
        "accent": "emerald",
        "header": re.compile(r"(?m)^\s*\(?b\)?\s*CODING\s+AND\s+DECODING\s+TEST\s*$", re.I),
    },
    {
        "id": "ch-07",
        "title": "Alphabetical Test",
        "slug": "alphabetical-test",
        "blurb": "Letter positions, order, and alphabet reasoning.",
        "accent": "violet",
        "header": re.compile(r"(?m)^\s*\(?c\)?\s*ALPHABETICAL\s+TEST\s*$", re.I),
    },
]

Q_START = re.compile(r"(?m)^\s*(\d{1,3})\s*[.)]\s+(.+)$")
BOOK_NOISE = re.compile(
    r"(?im)^(?:—\s*Book page.*|Dogar'?s Unique.*|Note:.*|Direction[s]?:.*|"
    r"What comes next.*|Write the correct words.*|Rearrange the following.*|"
    r"Point out the last letter.*|Insert the missing figures.*|"
    r"Write the missing figures.*|\*+|Cover also shows.*)$"
)
THEORY_SKIP = re.compile(
    r"(?i)dogar'?s unique|book page|cover also|transcription notes"
)


def read_source() -> tuple[str, list[str]]:
    doc = Document(str(SOURCE))
    parts: list[str] = []
    table_blobs: list[str] = []
    for p in doc.paragraphs:
        t = p.text.replace("\xa0", " ").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append(" | ".join(x for x in cells if x))
        table_blobs.append("\n".join(rows))
    return "\n".join(parts), table_blobs


def parse_key_map(text: str) -> dict[int, str]:
    keys: dict[int, str] = {}
    empty = set()
    for m in re.finditer(r"(?i)(\d{1,3})\s*[.)]\s*[—–\-]+\s*", text):
        empty.add(int(m.group(1)))
    for m in re.finditer(r"(?i)(\d{1,3})\s*[.)]\s*\(?\s*([a-d])\s*\)?", text):
        n = int(m.group(1))
        if n not in empty:
            keys[n] = m.group(2).upper()
    for m in re.finditer(r"(?im)^\s*(\d{1,3})\s*[.)]\s*\(([a-d])\)", text):
        n = int(m.group(1))
        if n not in empty:
            keys[n] = m.group(2).upper()
    for n in empty:
        keys.pop(n, None)
    return keys


def parse_options(block: str) -> list[str] | None:
    block = block.replace("\xa0", " ")
    block = re.sub(r"\(\s*6\s*\)", "(b)", block)
    block = re.sub(r"\(\s*([a-dA-D])\s*\)", lambda m: f"({m.group(1).lower()})", block)
    found: list[tuple[str, str]] = []
    for m in re.finditer(r"\(([a-d])\)\s*([^\n(]+)", block, re.I):
        letter = m.group(1).upper()
        val = re.sub(r"\s+", " ", m.group(2)).strip(" .,:;|\t")
        if val and len(val) <= 220:
            found.append((letter, val))
    if len(found) < 3:
        return None
    opts: dict[str, str] = {}
    next_missing = ["A", "B", "C", "D"]
    for letter, val in found:
        if letter not in opts:
            opts[letter] = val
            if letter in next_missing:
                next_missing.remove(letter)
        elif next_missing:
            opts[next_missing.pop(0)] = val
    if len(opts) < 4:
        return None
    return [opts[L] for L in "ABCD"]


def split_chapters(full_text: str) -> dict[str, str]:
    hits: list[tuple[int, str]] = []
    for ch in CHAPTER_DEFS:
        m = ch["header"].search(full_text)
        if m:
            hits.append((m.start(), ch["id"]))
    hits.sort(key=lambda x: x[0])
    sections: dict[str, str] = {}
    for i, (start, cid) in enumerate(hits):
        end = hits[i + 1][0] if i + 1 < len(hits) else len(full_text)
        # Trim at Word Building — beyond our MCQ scope
        wb = re.search(r"(?m)^\s*\(?d\)?\s*WORD\s+BUILDING\s+TEST", full_text[start:end], re.I)
        if wb:
            end = start + wb.start()
        if start < end:
            sections[cid] = full_text[start:end]
    return sections


def is_option_row(ln: str) -> bool:
    if re.match(r"(?i)^(hint|ans|explanation|solution|example|direction)\b", ln):
        return False
    if Q_START.match(ln):
        return False
    return bool(re.search(r"\([a-d]\)", ln, re.I))


def collect_option_lines(lines: list[str], start: int) -> tuple[list[str], int]:
    j = start
    opt_lines: list[str] = []
    while j < len(lines) and is_option_row(lines[j]):
        opt_lines.append(lines[j])
        j += 1
    return opt_lines, j


def theory_region(section: str) -> str:
    m = re.search(r"(?im)^\s*EXERCISE\b", section)
    if not m:
        return ""
    region = section[: m.start()]
    lines = region.splitlines()
    if lines:
        lines = lines[1:]  # drop chapter title
    return "\n".join(lines)


def _clean_stem(stem: str) -> str:
    stem = re.sub(r"\s+", " ", stem).strip(" .:-")
    stem = re.sub(r"\.+$", "", stem)
    if stem and not stem.endswith("?"):
        stem += "?"
    return stem


def extract_theory(section: str) -> dict:
    raw = theory_region(section)
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip() and not THEORY_SKIP.search(ln)]
    if not lines:
        return {"intro": "", "sections": []}

    intro_parts: list[str] = []
    sections: list[dict] = []
    current: dict | None = None
    cur_example: dict | None = None
    seen_heading = False

    def ensure_section(title: str = "Key concepts") -> dict:
        nonlocal current
        if current is None:
            current = {"title": title, "paragraphs": [], "examples": []}
        return current

    def flush_section() -> None:
        nonlocal current
        if current and (current["paragraphs"] or current["examples"]):
            sections.append(current)
        current = None

    i = 0
    while i < len(lines):
        line = lines[i]

        if re.match(r"^\([A-L]\)\s+", line):
            flush_section()
            seen_heading = True
            if "–" in line:
                title, rest = line.split("–", 1)
            elif " - " in line:
                title, rest = line.split(" - ", 1)
            else:
                title, rest = line, ""
            current = {
                "title": title.strip(),
                "paragraphs": [rest.strip()] if rest.strip() else [],
                "examples": [],
            }
            cur_example = None
            i += 1
            continue

        if re.match(r"(?i)^\([ivx]+\)\s*example", line):
            ensure_section("Coding patterns")
            stem = re.sub(r"(?i)^\([ivx]+\)\s*example\s*[\d:.]*\s*", "", line).strip()
            cur_example = {"stem": _clean_stem(stem), "options": [], "hint": "", "explanation": "", "answer": ""}
            current["examples"].append(cur_example)
            # Options may be on following lines
            opt_lines, j = collect_option_lines(lines, i + 1)
            if opt_lines:
                cur_example["options"] = parse_options("\n".join(opt_lines)) or []
                cur_example["stem"] = _clean_stem(cur_example["stem"])
                i = j
                continue
            i += 1
            continue

        if re.match(r"(?i)^example\s*[\d:.]*", line):
            ensure_section("Worked examples")
            stem = re.sub(r"(?i)^example\s*[\d:.]*[.:]?\s*", "", line).strip()
            cur_example = {"stem": _clean_stem(stem), "options": [], "hint": "", "explanation": "", "answer": ""}
            current["examples"].append(cur_example)
            if re.search(r"\([a-d]\)", line, re.I):
                cur_example["options"] = parse_options(line) or []
                cur_example["stem"] = _clean_stem(stem)
            else:
                opt_lines, j = collect_option_lines(lines, i + 1)
                if opt_lines:
                    cur_example["options"] = parse_options("\n".join(opt_lines)) or []
                    cur_example["stem"] = _clean_stem(stem)
                    i = j
                    continue
            i += 1
            continue

        if cur_example and not cur_example["options"] and re.search(r"\([a-d]\)", line, re.I):
            cur_example["options"] = parse_options(line) or []
            cur_example["stem"] = _clean_stem(cur_example["stem"])
            i += 1
            continue

        if re.match(r"(?i)^hint:", line):
            if cur_example:
                cur_example["hint"] = re.sub(r"(?i)^hint:\s*", "", line).strip()
            i += 1
            continue

        if re.match(r"(?i)^(explanation|solution)\.?\s*", line):
            text = re.sub(r"(?i)^(explanation|solution)\.?\s*", "", line).strip()
            if cur_example:
                cur_example["explanation"] = text
            else:
                ensure_section()["paragraphs"].append(text)
            i += 1
            continue

        ans_m = re.match(r"(?i)^ans\.?\s*\(?([a-d])\)?", line)
        if ans_m:
            if cur_example:
                cur_example["answer"] = ans_m.group(1).upper()
            cur_example = None
            i += 1
            continue

        if re.match(r"(?i)^direction\.?\s", line):
            i += 1
            continue

        if re.match(r"(?i)^in this type", line) and not seen_heading and not current:
            intro_parts.append(line)
            i += 1
            continue

        if not seen_heading and current is None and cur_example is None:
            intro_parts.append(line)
        else:
            ensure_section()["paragraphs"].append(line)
        i += 1

    flush_section()

    if not sections and intro_parts:
        sections = [{"title": "Overview", "paragraphs": [], "examples": []}]

    # Trim intro: stop before first example-like content absorbed elsewhere
    intro = " ".join(intro_parts)
    for split_at in ("Example:", "Example 1.", "Example 2.", "Example 3."):
        if split_at in intro:
            intro = intro.split(split_at)[0].strip()
            break

    # Drop empty examples and dedupe section titles
    cleaned_sections: list[dict] = []
    seen_titles: set[str] = set()
    for sec in sections:
        sec["examples"] = [
            ex for ex in sec["examples"]
            if ex.get("stem") and (ex.get("options") or ex.get("explanation"))
        ]
        title_key = sec["title"][:50].lower()
        if title_key in seen_titles and not sec["examples"]:
            # merge paragraphs into previous
            if cleaned_sections:
                cleaned_sections[-1]["paragraphs"].extend(sec["paragraphs"])
            continue
        seen_titles.add(title_key)
        if sec["paragraphs"] or sec["examples"]:
            cleaned_sections.append(sec)

    return {
        "intro": intro,
        "sections": cleaned_sections,
    }


def exercise_body(section: str) -> str:
    m = re.search(r"(?im)^\s*EXERCISE\b", section)
    body = section[m.start() :] if m else section
    cut = re.search(r"(?im)^\s*(?:Answers?\s*Key|Explanatory\s+Answers)\b", body)
    return body[: cut.start()] if cut else body


def answer_region(section: str) -> str:
    m = re.search(r"(?im)^\s*(?:Answers?\s*Key|Explanatory\s+Answers)\b", section)
    if not m:
        return ""
    region = section[m.start() :]
    stop = re.search(
        r"(?im)^\*{3,}|^\s*\(?[a-z]\)\s+(?:JUMBLED|CODING|ALPHABETICAL|WORD)\s+",
        region,
    )
    return region[: stop.start()] if stop and stop.start() > 0 else region


def extract_mcqs(body: str) -> list[dict]:
    lines = [ln.strip() for ln in body.splitlines() if ln.strip()]
    blocks: list[tuple[int, str]] = []
    cur_num = None
    cur_buf: list[str] = []

    def flush() -> None:
        nonlocal cur_num, cur_buf
        if cur_num is not None and cur_buf:
            blocks.append((cur_num, "\n".join(cur_buf)))
        cur_num, cur_buf = None, []

    for raw in lines:
        if BOOK_NOISE.match(raw):
            continue
        if re.match(r"(?i)^\d+\.\s*Analogy Test", raw):
            continue
        if re.match(r"(?i)^\(?[abc]\)?\s*(JUMBLED|CODING|ALPHABETICAL)", raw):
            continue
        m = Q_START.match(raw)
        if m and int(m.group(1)) <= 200 and len(m.group(2).strip()) >= 2:
            flush()
            cur_num = int(m.group(1))
            cur_buf = [m.group(2)]
            continue
        if cur_num is not None:
            if re.match(r"^\d{1,3}\s*[.)]\s*$", raw):
                flush()
                continue
            cur_buf.append(raw)
    flush()

    questions: list[dict] = []
    for num, block in blocks:
        options = parse_options(block)
        if not options:
            continue
        stem_part = re.split(r"\([aA]\)", block, maxsplit=1)[0]
        stem = re.sub(r"\s+", " ", stem_part).strip(" .:-")
        if len(stem) < 3:
            continue
        low = stem.lower()
        if low.startswith(("note:", "explanation", "ans.", "hint:", "example")):
            continue
        if not stem.endswith("?"):
            stem += "?"
        questions.append({"sourceNum": num, "question": stem[:500], "options": options})
    return questions


def jumbled_word_mcqs(body: str, word_answers: dict[int, str]) -> list[dict]:
    out: list[dict] = []
    pool = [w for w in word_answers.values() if len(w.replace(" ", "")) >= 3]
    for m in re.finditer(
        r"(?m)^\s*(\d{1,2})\s*[.)]\s+([A-Z]{3,20})(?:\s+\(([^)]+)\))?\s*$",
        body,
    ):
        num = int(m.group(1))
        jumble = m.group(2)
        hint = (m.group(3) or "").strip()
        ans_word = word_answers.get(num)
        if not ans_word:
            continue
        compact = re.sub(r"[^A-Za-z]", "", ans_word).upper()
        if len(compact) < 3:
            continue
        distractors: list[str] = []
        for w in pool:
            wc = re.sub(r"[^A-Za-z]", "", w).upper()
            if wc != compact and abs(len(wc) - len(compact)) <= 2 and wc not in distractors:
                distractors.append(wc)
            if len(distractors) >= 3:
                break
        letters = list(compact)
        j = 0
        while len(distractors) < 3 and j < len(letters) - 1:
            cand = letters[:]
            cand[j], cand[j + 1] = cand[j + 1], cand[j]
            w = "".join(cand)
            if w != compact and w not in distractors:
                distractors.append(w)
            j += 1
        options = [compact, *distractors[:3]]
        rot = num % 4
        options = options[rot:] + options[:rot]
        out.append({
            "sourceNum": num,
            "question": f"Unscramble '{jumble}'" + (f" ({hint})" if hint else "") + "?",
            "options": options,
            "answer": "ABCD"[options.index(compact)],
        })
    return out


def parse_jumbled_word_answers(answer_text: str) -> dict[int, str]:
    keys: dict[int, str] = {}
    for m in re.finditer(r"(?i)(\d{1,2})\s*[.)]\s*([A-Z][A-Z .'-]{1,24})", answer_text):
        word = m.group(2).strip(" .")
        if word.startswith("(") or re.fullmatch(r"[A-D]", word):
            continue
        keys[int(m.group(1))] = word.upper()
    return keys


def main() -> None:
    if not SOURCE.exists():
        raise SystemExit(f"Missing source file: {SOURCE}")

    full_text, table_blobs = read_source()
    sections = split_chapters(full_text)

    table_keys = [parse_key_map(blob) for blob in table_blobs]
    type1_keys = table_keys[0] if table_keys else {}
    type2_keys = table_keys[1] if len(table_keys) > 1 else {}
    type1_keys.setdefault(59, "B")
    type1_keys.setdefault(63, "B")
    type1_keys.setdefault(68, "D")
    chapter_override_keys = {"ch-01": type1_keys, "ch-02": type2_keys}

    OUT.mkdir(parents=True, exist_ok=True)
    for old in OUT.glob("ch-*.json"):
        old.unlink()
    for name in ("allQuestions.json", "chapters.json"):
        p = OUT / name
        if p.exists():
            p.unlink()

    collected: list[tuple[dict, list[dict], dict]] = []

    for cdef in CHAPTER_DEFS:
        section = sections.get(cdef["id"], "")
        if not section:
            continue
        theory = extract_theory(section)
        body = exercise_body(section)
        ans_text = answer_region(section)
        section_keys = parse_key_map(ans_text)
        override = chapter_override_keys.get(cdef["id"], {})

        qs = extract_mcqs(body)
        if cdef["id"] == "ch-05":
            word_answers = parse_jumbled_word_answers(ans_text)
            existing_nums = {q["sourceNum"] for q in qs}
            for jq in jumbled_word_mcqs(body, word_answers):
                if jq["sourceNum"] not in existing_nums:
                    qs.append(jq)

        seen: set[tuple] = set()
        uniq: list[dict] = []
        for q in qs:
            key = (q["sourceNum"], q["question"][:60].lower())
            if key in seen:
                continue
            seen.add(key)
            ans = q.get("answer") or override.get(q["sourceNum"]) or section_keys.get(q["sourceNum"])
            if not ans or any(o.startswith("Option ") for o in q["options"]):
                continue
            if re.search(r"(?i)decipher the following coded words", q["question"]):
                continue
            uniq.append({
                "sourceNum": q["sourceNum"],
                "question": q["question"],
                "options": q["options"],
                "answer": ans,
            })
        collected.append((cdef, uniq, theory))

    chapters_index = []
    flat = []
    qid = 1
    for cdef, qs, theory in collected:
        qs = sorted(qs, key=lambda x: x["sourceNum"])
        fixed = []
        for q in qs:
            item = {
                "id": qid,
                "sourceNum": q["sourceNum"],
                "question": q["question"],
                "options": q["options"],
                "answer": q["answer"],
                "answerVerified": True,
            }
            fixed.append(item)
            flat.append({
                "id": qid,
                "category": cdef["title"],
                "chapterId": cdef["id"],
                "question": item["question"],
                "options": item["options"],
                "answer": item["answer"],
                "answerVerified": True,
            })
            qid += 1

        if not fixed:
            continue

        payload = {
            "id": cdef["id"],
            "title": cdef["title"],
            "slug": cdef["slug"],
            "blurb": cdef["blurb"],
            "accent": cdef["accent"],
            "theory": theory,
            "questions": fixed,
        }
        (OUT / f"{cdef['id']}.json").write_text(
            json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
        )
        chapters_index.append({
            "id": cdef["id"],
            "title": cdef["title"],
            "slug": cdef["slug"],
            "blurb": cdef["blurb"],
            "accent": cdef["accent"],
            "questionCount": len(fixed),
            "verifiedCount": len(fixed),
            "figureBased": False,
            "keysOnly": False,
            "file": f"{cdef['id']}.json",
        })

    index = {
        "bookTitle": "IQ Practice Tests",
        "source": "Dogar_IQ_Verbal_Part.docx",
        "chapterCount": len(chapters_index),
        "totalQuestions": len(flat),
        "chapters": chapters_index,
    }
    (OUT / "chapters.json").write_text(json.dumps(index, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    (OUT / "allQuestions.json").write_text(json.dumps(flat, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    print(json.dumps({"chapters": len(chapters_index), "questions": len(flat)}, indent=2))
    for c in chapters_index:
        print(f"- {c['id']} {c['slug']}: {c['questionCount']} Qs · {c['title']}")


if __name__ == "__main__":
    main()
